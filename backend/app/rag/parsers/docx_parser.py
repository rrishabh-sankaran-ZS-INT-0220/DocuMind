from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from backend.app.rag.parsers.base import ParsedDocument, ParsedPage


class DOCXParser:
    async def parse(self, path: Path) -> ParsedDocument:
        doc = DocxDocument(path)
        text_lines: list[str] = []
        for para in doc.paragraphs:
            if para.text:
                text_lines.append(para.text)

        text = "\n".join(text_lines)
        page = ParsedPage(page_number=1, text=text, section_title=None)
        return ParsedDocument(pages=[page])
